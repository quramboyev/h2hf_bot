import logging
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application,
    CommandHandler,
    CallbackQueryHandler,
    MessageHandler,
    filters,
    ContextTypes,
    ConversationHandler,
    PicklePersistence
)
from telegram import ChatMember
from datetime import datetime, date, timedelta
import os
import sys
from dotenv import load_dotenv
import json
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ContextTypes
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, CallbackQueryHandler, ConversationHandler, ContextTypes
from typing import Dict, Any
from telegram.error import BadRequest, Forbidden

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# Получаем абсолютный путь к текущей директории
current_dir = os.path.dirname(os.path.abspath(__file__))
# Добавляем родительскую директорию в sys.path
sys.path.append(current_dir)

admins = [5012886318, 384221607, 491264374]

try:
    from bot5 import create_registration_doc
except ImportError as e:
    print(f"Error importing word_generator: {e}")
    print(f"Current directory: {current_dir}")
    print(f"Python path: {sys.path}")
    sys.exit(1)

# Настройка логирования
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.DEBUG
)
logger = logging.getLogger(__name__)

# Загрузка переменных окружения
load_dotenv()
TOKEN = os.getenv('6689906797:AAEKnJ1WtL7oiZGFnNd50tpHTsstXtcOvzM')
if not TOKEN:
    raise ValueError("No token provided! Add TELEGRAM_BOT_TOKEN to .env file.")

logger.info(f"Bot token loaded successfully")


# Функция для генерации документа со всеми регистрациями
async def generate_registrations_doc(tournament, registrations_list, lang):
    try:
        filename = create_registration_doc(
            registrations=registrations_list,
            tournament_data=tournament,
            language=lang
        )
        return filename
    except Exception as e:
        logger.error(f"Error generating Word document: {e}")
        return None


# Добавляем новое состояние
SUBSCRIPTION = 12

# Обновляем список состояний
SELECTING_LANGUAGE, SUBSCRIPTION, VIEWING_TOURNAMENTS, TOURNAMENT_DETAILS, DISCIPLINE_SELECT, REGION_SELECT, \
    INPUT_NAME, INPUT_BIRTH, GENDER_SELECT, WEIGHT_SELECT, INPUT_COACH, CONFIRM_DATA, EDIT_FIELD = range(13)

# Данные для переводов
translations = {
    'ru': {
        'select_language': 'Выберите язык:',
        'russian': '🇷🇺 Русский',
        'uzbek': '🇺🇿 Узбекский',
        'available_tournaments': 'Доступные турниры:',
        'tournament_info': '''
*{name}*
📅 Дата: {date}
🏢 Организатор: {organizer}
👥 Возраст: {min_age}-{max_age} лет
''',
        'register_button': '📝 Зарегистрироваться',
        'back_button': '« Назад',
        'select_discipline': 'Выберите дисциплину:',
        'self_defense': 'Самооборона',
        'combat': 'Поединок',
        'select_region': 'Выберите ваш регион из списка:',
        'enter_name': 'Введите ваше полное ФИО (Фамилия Имя Отчество) на латинице:',
        'enter_birth': 'Введите вашу дату рождения в формате ДД.ММ.ГГГГ\nПример: 15.03.2010',
        'select_gender': 'Выберите ваш пол:',
        'male': 'Мужчина',
        'female': 'Женщина',
        'select_weight': 'Выберите вашу весовую категорию:',
        'enter_coach': 'Введите полное ФИО вашего тренера (Фамилия Имя Отчество):',
        'invalid_date': 'Некорректная дата. Пожалуйста, введите дату в формате ДД.ММ.ГГГГ\nПример: 15.03.2011',
        'age_error': 'Ваш возраст не соответствует требованиям турнира (от {min_age} до {max_age} лет).\nРегистрация отменена.',
        'registration_closed': 'Регистрация на этот турнир закрыта. Она завершается за неделю до мероприятия.',
        'registration_summary': '''
*Проверьте введенные данные:*
🏆 Турнир: {tournament}
🎯 Дисциплина: {discipline}
🌍 Регион: {region}
👤 ФИО: {name}
📅 Дата рождения: {birth_date}
👫 Пол: {gender}
⚖️ Вес: {weight}
👨‍🏫 Тренер: {coach}
''',
        'confirm_correct': '✅ Всё верно',
        'restart': '🔄 Заново',
        'final_confirm': 'Вы подтверждаете регистрацию?',
        'yes': 'Да ✅',
        'no': 'Нет ❌',
        'registration_complete': 'Регистрация успешно завершена! ✅',
        'start_new': 'Для новой регистрации используйте /start',
        'error_message': 'Что-то пошло не так. Попробуйте еще раз.',
        'name_already_registered': 'ФИО "{name}" уже зарегистрировано в этой дисциплине. Пожалуйста, введите другое ФИО.',
        'subscribe_message': 'Пожалуйста, подпишитесь на наш канал, чтобы продолжить: https://t.me/fedrpo',
        'subscribe_button': 'Подписаться',
        'subscribe_confirm_button': 'Я подписался ✅',
    },
    'uz': {
        'select_language': 'Tilni tanlang:',
        'russian': '🇷🇺 Ruscha',
        'uzbek': '🇺🇿 O\'zbekcha',
        'available_tournaments': 'Mavjud musobaqalar:',
        'tournament_info': '''
*{name}*
📅 Sana: {date}
🏢 Tashkilotchi: {organizer}
👥 Yosh: {min_age}-{max_age} yosh
''',
        'register_button': '📝 Ro\'yxatdan o\'tish',
        'back_button': '« Orqaga',
        'select_discipline': 'Displinani tanlang:',
        'self_defense': 'O\'zini o\'zi himoya qilish',
        'combat': 'Jang',
        'select_region': 'Viloyatingizni ro\'yxatdan tanlang:',
        'enter_name': 'To\'liq F.I.O kiriting (Familiya Ism Sharif) lotincha:',
        'enter_birth': 'Tug\'ilgan sanangizni DD.MM.YYYY formatida kiriting\nMasalan: 15.03.2011',
        'select_gender': 'Jinsingizni tanlang:',
        'male': 'Erkak',
        'female': 'Ayol',
        'select_weight': 'Vazn toifangizni tanlang:',
        'enter_coach': 'Murabbiyingizning to\'liq F.I.O kiriting (Familiya Ism Sharif):',
        'invalid_date': 'Noto\'g\'ri sana. Iltimos, sanani DD.MM.YYYY formatida kiriting\nMasalan: 15.03.2010',
        'age_error': 'Yoshingiz musobaqa talablariga mos kelmaydi ({min_age} dan {max_age} yoshgacha).\nRo\'yxatdan o\'tish bekor qilindi.',
        'registration_closed': 'Ushbu musobaqaga ro\'yxatdan o\'tish yopilgan. Musobaqadan bir hafta oldin yopiladi.',
        'registration_summary': '''
*Ma\'lumotlarni tekshiring:*
🏆 Musobaqa: {tournament}
🎯 Disiplina: {discipline}
🌍 Viloyat: {region}
👤 F.I.O: {name}
📅 Tug\'ilgan sana: {birth_date}
👫 Jins: {gender}
⚖️ Vazn: {weight}
👨‍🏫 Murabbiy: {coach}
''',
        'confirm_correct': '✅ Hammasi to\'g\'ri',
        'restart': '🔄 Qayta boshlash',
        'final_confirm': 'Ro\'yxatdan o\'tishni tasdiqlaysizmi?',
        'yes': 'Ha ✅',
        'no': 'Yo\'q ❌',
        'registration_complete': 'Ro\'yxatdan o\'tish muvaffaqiyatli yakunlandi! ✅',
        'start_new': 'Yangi ro\'yxatdan o\'tish uchun /start buyrug\'idan foydalaning',
        'error_message': 'Xatolik yuz berdi. Iltimos, qaytadan urinib ko\'ring.',
        'name_already_registered': 'F.I.O "{name}" allaqachon ro\'yxatdan o\'tgan. Iltimos, boshqa F.I.O kiriting.',
        'subscribe_message': 'Iltimos, davom etish uchun kanalimizga obuna bo\'ling: https://t.me/fedrpo',
        'subscribe_button': 'Obuna bo\'lish',
        'subscribe_confirm_button': 'Men obuna bo\'ldim ✅',
    }
}

# Регионы
regions = [
    {"id": "toshkent viloyati", "nameRu": "Ташкентская облать", "nameUz": "Toshkent viloyati"},
    {"id": "toshkent shahri", "nameRu": "Город Ташкент", "nameUz": "Toshkent shahri"},
    {"id": "andijan", "nameRu": "Андижан", "nameUz": "Andijon"},
    {"id": "bukhara", "nameRu": "Бухара", "nameUz": "Buxoro"},
    {"id": "fergana", "nameRu": "Фергана", "nameUz": "Farg'ona"},
    {"id": "jizzakh", "nameRu": "Джизак", "nameUz": "Jizzax"},
    {"id": "namangan", "nameRu": "Наманган", "nameUz": "Namangan"},
    {"id": "navoiy", "nameRu": "Навои", "nameUz": "Navoiy"},
    {"id": "qashqadaryo", "nameRu": "Кашкадарья", "nameUz": "Qashqadaryo"},
    {"id": "samarkand", "nameRu": "Самарканд", "nameUz": "Samarqand"},
    {"id": "sirdaryo", "nameRu": "Сырдарья", "nameUz": "Sirdaryo"},
    {"id": "surxondaryo", "nameRu": "Сурхандарья", "nameUz": "Surxondaryo"},
    {"id": "xorazm", "nameRu": "Хорезм", "nameUz": "Xorazm"},
    {"id": "Qoraqalpog'ison", "nameRu": "Каракалпакстан", "nameUz": "Qoraqalpog'ison"},
    {"id": "IIV", "nameRu": "МВД", "nameUz": "IIV"},
    {"id": "MG", "nameRu": "НГ", "nameUz": "MG"},
    {"id": "MV", "nameRu": "МО", "nameUz": "MV"},
    {"id": "FVV", "nameRu": "МЧС", "nameUz": "FVV"},
    {"id": "DBQ", "nameRu": "ГТК", "nameUz": "DBQ"},
    {"id": "PDXX", "nameRu": "СГБП", "nameUz": "PDXX"},
    {"id": "DXX", "nameRu": "СГБ", "nameUz": "DXX"},
    {"id": "DXX ChQ", "nameRu": "ПВ СГБ", "nameUz": "DXX ChQ"},
    {"id": "IVV QQ", "nameRu": "КВ МВД", "nameUz": "IVV QQ"},
    {"id": "IIV Akademiyasi", "nameRu": "Академия МВД", "nameUz": "IIV Akademiyasi"},
    {"id": "MG JXU", "nameRu": "УОБ НГ", "nameUz": "MG JXU"},
    {"id": "MV QK Akademiyasi", "nameRu": "АВС МО", "nameUz": "MV QK Akademiyasi"},

]

# Весовые категории для мальчиков и девочек
weight_categories = {
    'male': ["21kg", "24kg", "27kg", "30kg", "33kg", "35kg", "38kg", "41kg", "44kg", "48kg", "52kg", "+52kg", "53kg",
             "57kg", "58kg", "62kg", "64kg", "67kg", "+67kg", "+70kg", "73kg", "80kg", "88kg", "97kg", "+97kg", ],
    'female': ["21kg", "24kg", "27kg", "30kg", "33kg", "35kg", "38kg", "+38kg", "48kg", "53kg", "58kg", "64kg", "70kg",
               "+70kg", ]
}

# Турниры
tournaments = [
    {
        "id": 1,
        "nameRu": "Чемпионат Узбекистана",
        "nameUz": "O'zbekiston Chempionati",
        "date": "08.04.2025",
        "organizerRu": "Федерация Рукопашного Боя Узбекистана",
        "organizerUz": "O'zbekiston Qo'l Jangi Federatsiyasi",
        "minAge": 10,
        "maxAge": 100
    },
    {
        "id": 2,
        "nameRu": "Кубок Ташкента",
        "nameUz": "Toshkent kubogi",
        "date": "20.06.2024",
        "organizerRu": "Федерация Рукопашного Боя Узбекистана",
        "organizerUz": "O'zbekiston Qo'l Jangi Federatsiyasi",
        "minAge": 14,
        "maxAge": 18
    }
]

# Список регистраций
registrations = []


# Функция для загрузки регистраций из файла
def load_registrations():
    global registrations
    if os.path.exists("registrations.json"):
        with open("registrations.json", "r", encoding="utf-8") as file:
            registrations = json.load(file)


# Функция для сохранения регистраций в файл
def save_registrations():
    with open("registrations.json", "w", encoding="utf-8") as file:
        json.dump(registrations, file, ensure_ascii=False, indent=4)


# Загружаем регистрации при старте
load_registrations()


def get_text(key: str, lang: str) -> str:
    """Получить перевод текста"""
    return translations.get(lang, {}).get(key, "Текст не найден")


def get_region_name(region_id: str, lang: str) -> str:
    """Получить название региона"""
    region = next((r for r in regions if r['id'] == region_id), None)
    return region[f'name{lang.capitalize()}'] if region else region_id


def is_registration_open(tournament_date: str) -> bool:
    """Проверка, открыта ли регистрация на турнир"""
    try:
        tournament_date = datetime.strptime(tournament_date, "%d.%m.%Y").date()
        return (tournament_date - date.today()) > timedelta(days=7)
    except Exception as e:
        logger.error(f"Error checking registration date: {e}")
        return False


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Начало разговора и выбор языка"""
    try:
        context.user_data.clear()  # Сбрасываем состояние
        keyboard = [
            [
                InlineKeyboardButton("🇷🇺 Русский", callback_data='lang_ru'),
                InlineKeyboardButton("🇺🇿 O'zbekcha", callback_data='lang_uz')
            ]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await update.message.reply_text(
            "Выберите язык / Tilni tanlang:",
            reply_markup=reply_markup
        )
        return SELECTING_LANGUAGE
    except Exception as e:
        logger.error(f"Error in start: {e}")
        await update.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def language_selected(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка выбора языка"""
    try:
        query = update.callback_query
        await query.answer()

        lang = query.data.split('_')[1]
        context.user_data['language'] = lang

        # Показываем сообщение с просьбой подписаться на канал
        keyboard = [
            [InlineKeyboardButton(get_text('subscribe_button', lang), url='https://t.me/fedrpo')],
            # Замените на ваш канал
            [InlineKeyboardButton(get_text('subscribe_confirm_button', lang), callback_data='subscribed')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(
            text=get_text('subscribe_message', lang),
            reply_markup=reply_markup
        )
        return SUBSCRIPTION
    except Exception as e:
        logger.error(f"Error in language_selected: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def handle_subscription(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка подтверждения подписки"""
    try:
        query = update.callback_query
        await query.answer()

        lang = context.user_data['language']

        # Показываем список турниров
        keyboard = []
        for tournament in tournaments:
            name = tournament[f'name{lang.capitalize()}']
            keyboard.append([InlineKeyboardButton(
                name, callback_data=f"tournament_{tournament['id']}"
            )])

        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(
            text=get_text('available_tournaments', lang),
            reply_markup=reply_markup
        )
        return VIEWING_TOURNAMENTS
    except Exception as e:
        logger.error(f"Error in handle_subscription: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def show_tournament_details(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Показать детали турнира"""
    try:
        query = update.callback_query
        await query.answer()

        lang = context.user_data['language']
        tournament_id = int(query.data.split('_')[1])
        tournament = next((t for t in tournaments if t['id'] == tournament_id), None)

        if not tournament:
            await query.edit_message_text("Tournament not found")
            return VIEWING_TOURNAMENTS

        context.user_data['tournament'] = tournament

        # Проверяем, открыта ли регистрация
        if not is_registration_open(tournament['date']):
            keyboard = [
                [InlineKeyboardButton(get_text('back_button', lang), callback_data="back_to_tournaments")]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(get_text('registration_closed', lang), reply_markup=reply_markup)
            return VIEWING_TOURNAMENTS  # Возвращаемся к состоянию просмотра турниров

        # Формируем информацию о турнире
        info = get_text('tournament_info', lang).format(
            name=tournament[f'name{lang.capitalize()}'],
            date=tournament['date'],
            organizer=tournament[f'organizer{lang.capitalize()}'],
            min_age=tournament['minAge'],
            max_age=tournament['maxAge']
        )

        keyboard = [
            [InlineKeyboardButton(get_text('register_button', lang), callback_data=f"register_{tournament['id']}")],
            [InlineKeyboardButton(get_text('back_button', lang), callback_data="back_to_tournaments")]
        ]

        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(
            text=info,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        return TOURNAMENT_DETAILS
    except Exception as e:
        logger.error(f"Error in show_tournament_details: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END

        # Формируем информацию о турнире
        info = get_text('tournament_info', lang).format(
            name=tournament[f'name{lang.capitalize()}'],
            date=tournament['date'],
            organizer=tournament[f'organizer{lang.capitalize()}'],
            min_age=tournament['minAge'],
            max_age=tournament['maxAge']
        )

        keyboard = [
            [InlineKeyboardButton(get_text('register_button', lang), callback_data=f"register_{tournament['id']}")],
            [InlineKeyboardButton(get_text('back_button', lang), callback_data='back_to_tournaments')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(
            text=info,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        return TOURNAMENT_DETAILS
    except Exception as e:
        logger.error(f"Error in show_tournament_details: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def start_registration(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Начало процесса регистрации"""
    try:
        query = update.callback_query
        await query.answer()

        lang = context.user_data['language']

        # Показываем кнопки с дисциплинами
        keyboard = [
            [InlineKeyboardButton(get_text('self_defense', lang), callback_data='discipline_self_defense')],
            [InlineKeyboardButton(get_text('combat', lang), callback_data='discipline_combat')],
            [InlineKeyboardButton(get_text('back_button', lang), callback_data='back_to_tournament')]
        ]

        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(
            text=get_text('select_discipline', lang),
            reply_markup=reply_markup
        )
        return DISCIPLINE_SELECT
    except Exception as e:
        logger.error(f"Error in start_registration: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def discipline_selected(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка выбора дисциплины"""
    try:
        query = update.callback_query
        await query.answer()

        discipline = query.data.split('_')[1]
        context.user_data['discipline'] = discipline

        # Логируем выбранную дисциплину и текущие данные пользователя
        logger.info(f"Selected discipline: {discipline}")
        logger.info(f"Current user_data: {context.user_data}")

        lang = context.user_data['language']

        # Показываем кнопки с регионами
        keyboard = []
        row = []
        for i, region in enumerate(regions):
            row.append(InlineKeyboardButton(
                region[f'name{lang.capitalize()}'],
                callback_data=f"region_{region['id']}"
            ))
            if len(row) == 2 or i == len(regions) - 1:
                keyboard.append(row)
                row = []

        keyboard.append([InlineKeyboardButton(
            get_text('back_button', lang),
            callback_data="back_to_discipline"
        )])

        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(
            text=get_text('select_region', lang),
            reply_markup=reply_markup
        )
        return REGION_SELECT
    except Exception as e:
        logger.error(f"Error in discipline_selected: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def region_selected(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка выбора региона"""
    try:
        query = update.callback_query
        await query.answer()

        region_id = query.data.split('_')[1]
        context.user_data['region'] = region_id

        lang = context.user_data['language']
        await query.edit_message_text(get_text('enter_name', lang))
        return INPUT_NAME
    except Exception as e:
        logger.error(f"Error in region_selected: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def process_name(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка ввода имени"""
    try:
        name = update.message.text.strip()
        # Проверяем, что ФИО содержит хотя бы два слова
        if len(name.split()) < 2:
            lang = context.user_data['language']
            await update.message.reply_text(get_text('enter_name', lang))
            return INPUT_NAME

        # Проверяем, что ФИО не используется повторно в одной дисциплине
        tournament_id = context.user_data['tournament']['id']
        discipline = context.user_data['discipline']
        for reg in registrations:
            if (reg['tournamentId'] == tournament_id and
                    reg['discipline'] == discipline and
                    reg['fullName'].lower() == name.lower()):
                lang = context.user_data['language']
                await update.message.reply_text(
                    get_text('name_already_registered', lang).format(name=name))
                return INPUT_NAME

        context.user_data['name'] = name
        lang = context.user_data['language']

        await update.message.reply_text(get_text('enter_birth', lang))
        return INPUT_BIRTH
    except Exception as e:
        logger.error(f"Error in process_name: {e}")
        await update.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def process_birth_date(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка ввода даты рождения"""
    try:
        lang = context.user_data['language']

        # Парсим дату
        birth_date = datetime.strptime(update.message.text, "%d.%m.%Y").date()
        tournament = context.user_data['tournament']

        # Проверяем возраст
        age = (date.today() - birth_date).days // 365
        if age < tournament['minAge'] or age > tournament['maxAge']:
            # Сообщение об ошибке возраста
            error_message = get_text('age_error', lang).format(
                min_age=tournament['minAge'],
                max_age=tournament['maxAge']
            )

            # Добавляем кнопку "Назад"
            keyboard = [
                [InlineKeyboardButton(get_text('back_button', lang), callback_data="back_to_tournaments")]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)

            await update.message.reply_text(
                error_message,
                reply_markup=reply_markup
            )
            return VIEWING_TOURNAMENTS  # Возвращаемся к списку турниров
        else:
            # Если возраст подходит, продолжаем регистрацию
            context.user_data['birth_date'] = update.message.text

            # Показываем кнопки с выбором пола
            keyboard = [
                [InlineKeyboardButton(get_text('male', lang), callback_data='gender_male')],
                [InlineKeyboardButton(get_text('female', lang), callback_data='gender_female')]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await update.message.reply_text(
                get_text('select_gender', lang),
                reply_markup=reply_markup
            )
            return GENDER_SELECT
    except ValueError:
        await update.message.reply_text(get_text('invalid_date', lang))
        return INPUT_BIRTH
    except Exception as e:
        logger.error(f"Error in process_birth_date: {e}")
        await update.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def gender_selected(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка выбора пола"""
    try:
        query = update.callback_query
        await query.answer()

        gender = query.data.split('_')[1]
        context.user_data['gender'] = gender

        lang = context.user_data['language']

        # Показываем кнопки с весовыми категориями в зависимости от пола
        weights = weight_categories[gender]
        keyboard = []
        row = []
        for i, weight in enumerate(weights):
            row.append(InlineKeyboardButton(
                weight,
                callback_data=f"weight_{weight}"
            ))
            if len(row) == 3 or i == len(weights) - 1:
                keyboard.append(row)
                row = []

        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(
            get_text('select_weight', lang),
            reply_markup=reply_markup
        )
        return WEIGHT_SELECT
    except Exception as e:
        logger.error(f"Error in gender_selected: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def weight_selected(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка выбора весовой категории"""
    try:
        query = update.callback_query
        await query.answer()

        weight = query.data.split('_')[1]
        context.user_data['weight'] = weight

        lang = context.user_data['language']
        await query.edit_message_text(get_text('enter_coach', lang))
        return INPUT_COACH
    except Exception as e:
        logger.error(f"Error in weight_selected: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def process_coach(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка ввода имени тренера"""
    try:
        coach = update.message.text.strip()
        # Проверяем, что ФИО тренера содержит хотя бы два слова
        if len(coach.split()) < 2:
            lang = context.user_data['language']
            await update.message.reply_text(get_text('enter_coach', lang))
            return INPUT_COACH

        context.user_data['coach'] = coach
        lang = context.user_data['language']

        # Логируем текущие данные пользователя
        logger.info(f"Current user_data before summary: {context.user_data}")

        # Получаем название дисциплины
        discipline_key = context.user_data['discipline']
        discipline_text = get_text(discipline_key, lang)

        # Логируем полученный текст дисциплины
        logger.info(f"Discipline key: {discipline_key}, Discipline text: {discipline_text}")

        # Формируем сводку
        tournament = context.user_data['tournament']
        summary = get_text('registration_summary', lang).format(
            tournament=tournament[f'name{lang.capitalize()}'],
            discipline=discipline_text,  # Используем полученный текст
            region=get_region_name(context.user_data['region'], lang),
            name=context.user_data['name'],
            birth_date=context.user_data['birth_date'],
            gender=get_text(context.user_data['gender'], lang),
            weight=context.user_data['weight'],
            coach=context.user_data['coach']
        )

        keyboard = [
            [
                InlineKeyboardButton(get_text('confirm_correct', lang),
                                     callback_data='confirm_correct'),
                InlineKeyboardButton(get_text('restart', lang),
                                     callback_data='restart')
            ]
        ]

        reply_markup = InlineKeyboardMarkup(keyboard)
        await update.message.reply_text(
            summary,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        return CONFIRM_DATA
    except Exception as e:
        logger.error(f"Error in process_coach: {e}")
        await update.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def handle_confirmation(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка подтверждения данных"""
    try:
        query = update.callback_query
        await query.answer()

        lang = context.user_data['language']

        if query.data == 'restart':
            # Начинаем регистрацию заново
            context.user_data.clear()  # Сбрасываем состояние
            await query.edit_message_text(get_text('start_new', lang))
            return await start(update, context)
        else:
            # Показываем финальное подтверждение
            keyboard = [
                [
                    InlineKeyboardButton(get_text('yes', lang),
                                         callback_data='final_confirm_yes'),
                    InlineKeyboardButton(get_text('no', lang),
                                         callback_data='final_confirm_no')
                ]
            ]

            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(
                get_text('final_confirm', lang),
                reply_markup=reply_markup
            )
            return CONFIRM_DATA
    except Exception as e:
        logger.error(f"Error in handle_confirmation: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def handle_final_confirmation(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка финального подтверждения"""
    try:
        query = update.callback_query
        await query.answer()

        lang = context.user_data['language']

        if query.data == 'final_confirm_yes':
            # Сохраняем регистрацию
            registration = {
                "tournamentId": context.user_data['tournament']['id'],
                "discipline": context.user_data['discipline'],
                "region": context.user_data['region'],
                "fullName": context.user_data['name'],
                "birthDate": context.user_data['birth_date'],
                "gender": context.user_data['gender'],
                "weightCategory": context.user_data['weight'],
                "coachName": context.user_data['coach'],
                "createdAt": datetime.now().isoformat()
            }
            registrations.append(registration)
            save_registrations()  # Сохраняем регистрации в файл

            # Генерация документа для администратора (без отправки пользователю)
            generated_doc = await generate_registrations_doc(context.user_data['tournament'], registrations, lang)
            if not generated_doc:
                logger.error("Failed to generate registration document")
            else:
                logger.info(f"Document generated successfully: {generated_doc}")

            # Отправляем только подтверждение пользователю
            await query.edit_message_text(get_text('registration_complete', lang))
            await query.message.reply_text(get_text('start_new', lang))
            return ConversationHandler.END
        else:
            # Возвращаемся к сводке
            return await process_coach(update, context)
    except Exception as e:
        logger.error(f"Error in handle_final_confirmation: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def back_to_tournaments(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Возврат к списку турниров"""
    try:
        query = update.callback_query
        await query.answer()

        lang = context.user_data.get('language', 'ru')

        # Показываем список турниров
        keyboard = []
        for tournament in tournaments:
            name = tournament[f'name{lang.capitalize()}']
            keyboard.append([InlineKeyboardButton(
                name, callback_data=f"tournament_{tournament['id']}"
            )])

        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(
            text=get_text('available_tournaments', lang),
            reply_markup=reply_markup
        )
        return VIEWING_TOURNAMENTS
    except Exception as e:
        logger.error(f"Error in back_to_tournaments: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def back_to_tournament(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Возврат к деталям турнира"""
    try:
        query = update.callback_query
        await query.answer()
        logger.info("Back to tournament button pressed")  # Логирование

        lang = context.user_data.get('language', 'ru')
        tournament = context.user_data.get('tournament')

        if not tournament:
            await query.edit_message_text(get_text('error_message', lang))
            return ConversationHandler.END

        # Формируем информацию о турнире
        info = get_text('tournament_info', lang).format(
            name=tournament[f'name{lang.capitalize()}'],
            date=tournament['date'],
            organizer=tournament[f'organizer{lang.capitalize()}'],
            min_age=tournament['minAge'],
            max_age=tournament['maxAge']
        )

        keyboard = [
            [InlineKeyboardButton(get_text('register_button', lang), callback_data=f"register_{tournament['id']}")],
            [InlineKeyboardButton(get_text('back_button', lang), callback_data="back_to_tournaments")]
        ]

        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(
            text=info,
            reply_markup=reply_markup,
            parse_mode='Markdown'
        )
        return TOURNAMENT_DETAILS
    except Exception as e:
        logger.error(f"Error in back_to_tournament: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def back_to_discipline(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Возврат к выбору дисциплины"""
    try:
        query = update.callback_query
        await query.answer()
        logger.info("Back to discipline button pressed")  # Логирование

        lang = context.user_data.get('language', 'ru')

        # Показываем кнопки с дисциплинами
        keyboard = [
            [InlineKeyboardButton(get_text('self_defense', lang), callback_data='discipline_self_defense')],
            [InlineKeyboardButton(get_text('combat', lang), callback_data='discipline_combat')],
            [InlineKeyboardButton(get_text('back_button', lang), callback_data='back_to_tournament')]
        ]

        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(
            text=get_text('select_discipline', lang),
            reply_markup=reply_markup
        )
        return DISCIPLINE_SELECT
    except Exception as e:
        logger.error(f"Error in back_to_discipline: {e}")
        await update.callback_query.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def clear_command(update:Update, context: ContextTypes.DEFAULT_TYPE):
    # Проверяем, является ли пользователь администратором
    if update.effective_user.id not in admins:
        await update.message.reply_text("У вас нет прав для выполнения этой команды.")
        return
    os.remove('registrations.json')
    await update.message.reply_text('registration.json was cleared')



async def download_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Создание и отправка документа со списком регистраций"""
    try:
        # Проверяем, является ли пользователь администратором
        if update.effective_user.id not in admins:
            await update.message.reply_text("У вас нет прав для выполнения этой команды.")
            return

        # Загружаем регистрации из файла
        if not os.path.exists("registrations.json"):
            await update.message.reply_text("Нет данных о регистрациях.")
            return

        with open("registrations.json", "r", encoding="utf-8") as file:
            registrations_data = json.load(file)

        # Находим соответствующий турнир
        tournament_data = next((t for t in tournaments if t['id'] == 1), None)
        if not tournament_data:
            await update.message.reply_text("Турнир не найден.")
            return

        # Генерируем документ
        filename = create_doc(tournament_data, registrations_data, 'ru')
        
        if not filename:
            await update.message.reply_text("Ошибка при создании документа.")
            return

        # Отправляем документ
        with open(filename, 'rb') as doc:
            await update.message.reply_document(
                document=doc,
                filename=f"Регистрации_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
            )
        
        # Удаляем временный файл
        os.remove(filename)

    except Exception as e:
        logger.error(f"Error in download_command: {e}")
        await update.message.reply_text("Произошла ошибка при создании документа.")


def create_doc(registrations, tournament_data, language='ru'):
    """
    Создает Word документ с регистрациями
    Args:
        registrations: список регистраций
        tournament_data: данные турнира
        language: язык документа ('ru' или 'uz')
    Returns:
        str: путь к созданному файлу
    """
    try:
        doc = Document()

        # Настройка полей страницы
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # Заголовок документа
        title = doc.add_paragraph()
        title_run = title.add_run(f"Список участников - {tournament_data[f'name{language.capitalize()}']}")
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Добавляем информацию о турнире
        info = doc.add_paragraph()
        info.add_run(f"Дата турнира: {tournament_data['date']}\n")
        info.add_run(f"Организатор: {tournament_data[f'organizer{language.capitalize()}']}")
        info.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Создаем таблицы для каждой дисциплины
        disciplines = {
            'self_defense': 'Самооборона',
            'combat': 'Поединок'
        }

        for disc_key, disc_name in disciplines.items():
            # Фильтруем регистрации по дисциплине
            disc_registrations = [r for r in registrations if r['discipline'] == disc_key]

            if disc_registrations:
                # Добавляем заголовок дисциплины
                doc.add_paragraph().add_run(f"\n{disc_name}").bold = True

                # Создаем таблицу
                table = doc.add_table(rows=1, cols=7)
                table.style = 'Table Grid'

                # Заголовки таблицы
                headers = ['№', 'ФИО', 'Дата рождения', 'Регион', 'Пол', 'Вес', 'Тренер']
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Заполняем таблицу данными
                for idx, reg in enumerate(disc_registrations, 1):
                    row_cells = table.add_row().cells

                    # Номер
                    row_cells[0].text = str(idx)
                    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # ФИО
                    row_cells[1].text = reg['fullName']

                    # Дата рождения
                    row_cells[2].text = reg['birthDate']
                    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Регион
                    region = next((r for r in regions if r['id'] == reg['region']), None)
                    row_cells[3].text = region[f'name{language.capitalize()}'] if region else reg['region']

                    # Пол
                    row_cells[4].text = 'Мужской' if reg['gender'] == 'male' else 'Женский'
                    row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Вес
                    row_cells[5].text = reg['weightCategory']
                    row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Тренер
                    row_cells[6].text = reg['coachName']

                # Устанавливаем ширину столбцов
                widths = [1, 4, 2, 3, 1.5, 1.5, 4]
                for i, width in enumerate(widths):
                    for cell in table.columns[i].cells:
                        cell.width = Cm(width)

        # Сохраняем документ
        filename = f"registrations_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        doc.save(filename)
        return filename

    except Exception as e:
        return None


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Отмена регистрации"""
    try:
        lang = context.user_data.get('language', 'ru')
        await update.message.reply_text(get_text('start_new', lang))
        context.user_data.clear()  # Сбрасываем состояние
        return ConversationHandler.END
    except Exception as e:
        logger.error(f"Error in cancel: {e}")
        await update.message.reply_text(get_text('error_message', 'ru'))
        return ConversationHandler.END


async def error_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик ошибок"""
    logger.error(f"Exception while handling an update: {context.error}")
    try:
        if "Conflict: terminated by other getUpdates request" in str(context.error):
            logger.warning("Detected multiple bot instances, attempting to recover...")
            return

        if update and update.effective_message:
            await update.effective_message.reply_text(
                "Произошла ошибка при обработке запроса. Пожалуйста, попробуйте снова /start"
            )
    except Exception as e:
        logger.error(f"Error in error handler: {e}")


def main():
    """Запуск бота"""
    try:
        # Удаляем файл bot_data при каждом запуске
        if os.path.exists("bot_data"):
            os.remove("bot_data")

        # Инициализируем приложение
        application = Application.builder() \
            .token(TOKEN) \
            .persistence(PicklePersistence(filepath="bot_data")) \
            .build()

        # Создаем обработчик разговора
        conv_handler = ConversationHandler(
            entry_points=[CommandHandler('start', start)],
            states={
                SELECTING_LANGUAGE: [
                    CallbackQueryHandler(language_selected, pattern='^lang_')
                ],
                SUBSCRIPTION: [
                    CallbackQueryHandler(handle_subscription, pattern='^subscribed$')
                ],
                VIEWING_TOURNAMENTS: [
                    CallbackQueryHandler(show_tournament_details, pattern='^tournament_'),
                    CallbackQueryHandler(back_to_tournaments, pattern='^back_to_tournaments$')
                ],
                TOURNAMENT_DETAILS: [
                    CallbackQueryHandler(start_registration, pattern='^register_'),
                    CallbackQueryHandler(back_to_tournaments, pattern='^back_to_tournaments$')
                ],
                DISCIPLINE_SELECT: [
                    CallbackQueryHandler(discipline_selected, pattern='^discipline_'),
                    CallbackQueryHandler(back_to_tournament, pattern='^back_to_tournament$')
                ],
                REGION_SELECT: [
                    CallbackQueryHandler(region_selected, pattern='^region_'),
                    CallbackQueryHandler(back_to_discipline, pattern='^back_to_discipline$')
                ],
                INPUT_NAME: [
                    MessageHandler(filters.TEXT & ~filters.COMMAND, process_name)
                ],
                INPUT_BIRTH: [
                    MessageHandler(filters.TEXT & ~filters.COMMAND, process_birth_date)
                ],
                GENDER_SELECT: [
                    CallbackQueryHandler(gender_selected, pattern='^gender_')
                ],
                WEIGHT_SELECT: [
                    CallbackQueryHandler(weight_selected, pattern='^weight_')
                ],
                INPUT_COACH: [
                    MessageHandler(filters.TEXT & ~filters.COMMAND, process_coach)
                ],
                CONFIRM_DATA: [
                    CallbackQueryHandler(handle_confirmation, pattern='^confirm_'),
                    CallbackQueryHandler(handle_final_confirmation, pattern='^final_confirm_')
                ],
            },
            fallbacks=[
                CommandHandler('start', start),  # /start как fallback
                CommandHandler('cancel', cancel)
            ],
            name="tournament_registration",
            persistent=True
        )

        # Добавляем обработчики
        application.add_handler(conv_handler)
        application.add_error_handler(error_handler)
        application.add_handler(CommandHandler('download', download_command))
        application.add_handler(CommandHandler('clear', clear_command))

        # Запускаем бота
        logger.info("Starting bot...")
        application.run_polling(allowed_updates=Update.ALL_TYPES)
    except Exception as e:
        logger.error(f"Error in main: {e}")


if __name__ == '__main__':
    main()
