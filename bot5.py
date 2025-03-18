from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from datetime import datetime
import os

def create_registration_doc(registrations, tournament_data, language='ru'):
    """
    Создаёт Word-документ со всеми регистрациями для одного турнира.
    Сортирует данные сначала по регионам, затем по весу (от меньшего к большему).
    """
    filename = f"registrations_{tournament_data['id']}.docx"  # Файл для турнира
    # Создаем новый документ
    doc = Document()
    title = tournament_data[f'name{language.capitalize()}']
    doc.add_heading(title, 0)
    doc.add_paragraph(f"{'Дата' if language == 'ru' else 'Sana'}: {tournament_data['date']}")
    doc.add_paragraph(f"{'Организатор' if language == 'ru' else 'Tashkilotchi'}: {tournament_data[f'organizer{language.capitalize()}']}")

    # Создаём таблицу заголовков
    table = doc.add_table(rows=1, cols=8)  # Добавляем колонку для дисциплины и пола
    table.style = 'Table Grid'
    headers = {
        'ru': ['№', 'ФИО', 'Регион', 'Дата рождения', 'Пол', 'Вес', 'Тренер', 'Дисциплина'],
        'uz': ['№', 'F.I.O', 'Viloyat', "Tug'ilgan sana", 'Jins', 'Vazn', 'Murabbiy', 'Disiplina']
    }
    for i, header in enumerate(headers[language]):
        table.rows[0].cells[i].text = header

    # Функция для преобразования веса в число для сортировки
    def parse_weight(weight):
        """Преобразует вес в число для корректной сортировки."""
        if weight.endswith('kg'):
            weight = weight[:-2]  # Убираем 'kg'
        if '+' in weight:
            return int(weight[:-1]) + 0.5  # Вес с '+' будет больше
        return int(weight)

    # Сортировка по региону -> весу (от меньшего к большему)
    sorted_registrations = sorted(
        registrations,
        key=lambda r: (
            r['region'],  # Сначала сортируем по региону
            parse_weight(r['weightCategory'])  # Затем по весу
        )
    )

    # Добавляем данные в таблицу
    for idx, reg in enumerate(sorted_registrations, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = reg['fullName']
        row_cells[2].text = reg['region']
        row_cells[3].text = reg['birthDate']
        row_cells[4].text = 'Erkak' if reg['gender'] == 'male' else 'Ayol'
        row_cells[5].text = reg['weightCategory']
        row_cells[6].text = reg['coachName']
        row_cells[7].text = 'Jang' if reg['discipline'] == 'combat' else 'Self-defence'

    # Форматируем таблицу
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    # Сохраняем файл
    doc.save(filename)
    return filename
